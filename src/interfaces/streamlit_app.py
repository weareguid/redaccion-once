# Aplicación Streamlit para Once Noticias - Sistema Optimizado
import os
import sys

# Asegurar que el directorio raíz está en el path
if os.path.abspath('.') not in sys.path:
    sys.path.insert(0, os.path.abspath('.'))

import streamlit as st
import openai
from datetime import datetime
from typing import Dict, Any
import json
import time

# Librerías para exportación
try:
    from docx import Document
    from docx.shared import Inches
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    import io
    EXPORT_AVAILABLE = True
except ImportError:
    EXPORT_AVAILABLE = False

# Importaciones del sistema optimizado
from src.core.prompt_system import OptimizedOnceNoticiasPromptSystem
from src.core.quality_assurance import OnceNoticiasQualityAssurance
from config.settings import config

# Configuración de la página
st.set_page_config(
    page_title=config.STREAMLIT_CONFIG["page_title"],
    page_icon=config.STREAMLIT_CONFIG["page_icon"],
    layout=config.STREAMLIT_CONFIG["layout"],
    initial_sidebar_state="collapsed"  # Sidebar sin contenido
)

def get_web_search_setting():
    """Determina la configuración de force_web_search basada en la selección actual"""
    web_search_mode = st.session_state.get('web_search_mode', '🤖 Auto (recomendado)')

    if web_search_mode == "🌐 Siempre activada":
        return True
    elif web_search_mode == "📚 Solo conocimiento base":
        return False
    else:  # Auto (recomendado)
        return None

def initialize_system():
    """Inicializa el sistema editorial optimizado"""

    # Configurar OpenAI
    openai_api_key = config.get_api_key("OPENAI_API_KEY")
    if not openai_api_key:
        st.error("❌ OPENAI_API_KEY no configurada. Configúrala en config/.streamlit/secrets.toml")
        st.stop()

    openai.api_key = openai_api_key
    client = openai.OpenAI(api_key=openai_api_key)

    # Inicializar sistemas (sin forzar base de datos)
    prompt_system = OptimizedOnceNoticiasPromptSystem(client, enable_database=None)
    qa_system = OnceNoticiasQualityAssurance()

    return prompt_system, qa_system

def save_new_content_metrics(prompt_system, content_data, user_rating=None):
    """
    Guarda métricas para contenido NUEVO generado por LLM - crea nuevo registro

    Returns:
        int: ID del registro creado, o None si falló
    """
    # Preparar datos completos para guardado
    complete_data = {
        **content_data,
        'user_rating': user_rating,
        'sources_prompt': content_data.get('user_sources', ''),
        'generated_content': content_data.get('content', ''),
        'word_count': len(content_data.get('content', '').split()) if content_data.get('content') else 0,
        'improvement_applied': content_data.get('improvement_applied', False),
        'web_search_used': content_data.get('web_search_used', False),
        'web_search_sources_count': len(content_data.get('citations', [])),
        'citations_data': json.dumps(content_data.get('citations', []), ensure_ascii=False),
        'generation_time_seconds': content_data.get('generation_time', 0),
        'tokens_used': content_data.get('token_count', 0),
        'openai_call_id': content_data.get('openai_call_id', '')  # ✅ Add OpenAI call ID
    }

    # Guardar usando el sistema de prompt - esto devuelve el ID del registro
    return prompt_system.save_to_database(complete_data)

def update_content_rating_only(record_id, user_rating, user_feedback=None):
    """
    Actualiza solo el rating de un contenido existente - NO crea nuevo registro
    """
    from src.utils.database_connection import update_content_rating
    return update_content_rating(record_id, user_rating, user_feedback)

def update_content_rating_by_openai_id(openai_call_id, user_rating):
    """
    Actualiza el rating de un contenido existente usando OpenAI call ID
    ✅ MULTI-USER SAFE: Uses OpenAI call ID instead of record ID
    """
    from src.utils.database_connection import update_content_rating_by_openai_id
    return update_content_rating_by_openai_id(openai_call_id, user_rating)

def display_content_with_citations(content: str, citations: list, user_sources: str = ""):
    """
    Muestra contenido con citaciones inline clickeables y fuentes del usuario
    """
    if not citations and not user_sources:
        st.markdown(content)
        return

    # Mostrar contenido principal
    st.markdown(content)

    # Mostrar fuentes si hay citaciones o fuentes del usuario
    if citations or user_sources:
        st.markdown("---")
        st.markdown("### 📚 Fuentes consultadas:")

        citation_count = 0

        # Mostrar citaciones de web search si las hay
        if citations:
            # Agrupar citaciones únicas por URL
            unique_citations = {}
            for citation in citations:
                url = citation.get('url', '')

                if url and url not in unique_citations:
                    # Actualizar citation con URL limpia
                    citation_copy = citation.copy()
                    citation_copy['url'] = url
                    unique_citations[url] = citation_copy

            # Mostrar cada fuente con formato mejorado
            for i, (url, citation) in enumerate(unique_citations.items(), 1):
                title = citation.get('title', 'Fuente sin título')

                # Acortar títulos muy largos
                if len(title) > 80:
                    title = title[:77] + "..."

                # Crear tarjeta de citación
                st.markdown(f"""
                **{i}.** [{title}]({url})
                """)
                citation_count = i

        # Mostrar fuentes del usuario si las hay
        if user_sources and user_sources.strip():
            st.markdown("#### 📝 Fuentes proporcionadas por el usuario:")

            # Procesar fuentes del usuario línea por línea
            sources_lines = [line.strip() for line in user_sources.split('\n') if line.strip()]
            for i, source in enumerate(sources_lines, citation_count + 1):
                # Detectar si es un URL
                if source.startswith('http'):
                    st.markdown(f"**{i}.** [{source}]({source})")
                else:
                    st.markdown(f"**{i}.** {source}")

        # Mensaje final
        total_sources = len(citations) if citations else 0
        user_sources_count = len([line for line in user_sources.split('\n') if line.strip()]) if user_sources else 0

        if total_sources > 0 and user_sources_count > 0:
            st.markdown(f"*{total_sources} fuente(s) automática(s) + {user_sources_count} del usuario*")
        elif total_sources > 0:
            st.markdown(f"*{total_sources} fuente(s) consultada(s) automáticamente*")
        elif user_sources_count > 0:
            st.markdown(f"*{user_sources_count} fuente(s) proporcionada(s) por el usuario*")

def create_word_document(content: str, metadata: Dict) -> io.BytesIO:
    """Crea un documento Word con el contenido generado"""
    if not EXPORT_AVAILABLE:
        return None

    doc = Document()

    # Agregar encabezado Once Noticias
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "Once Noticias - Sistema Editorial Optimizado"

    # Título del documento
    title = doc.add_heading('Once Noticias', 0)
    title.alignment = 1  # Centrado

    # Metadatos
    doc.add_heading('Información del Contenido', level=2)
    info_para = doc.add_paragraph()
    info_para.add_run(f"Categoría: ").bold = True
    info_para.add_run(f"{metadata.get('category', '')} / {metadata.get('subcategory', '')}\n")
    info_para.add_run(f"Tipo: ").bold = True
    info_para.add_run(f"{metadata.get('text_type', '')}\n")
    info_para.add_run(f"Fecha: ").bold = True
    info_para.add_run(f"{datetime.now().strftime('%d/%m/%Y %H:%M')}\n")
    info_para.add_run(f"Tokens: ").bold = True
    info_para.add_run(f"{metadata.get('token_count', 0):,}")

    # Separador
    doc.add_page_break()

    # Contenido principal
    doc.add_heading('Contenido Generado', level=1)

    # Agregar contenido párrafo por párrafo
    for paragraph in content.split('\n\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())

    # Pie de página
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Generado por Once Noticias AI - {datetime.now().strftime('%d/%m/%Y')}"

    # Guardar en BytesIO
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def create_pdf_document(content: str, metadata: Dict) -> io.BytesIO:
    """Crea un documento PDF con el contenido generado"""
    if not EXPORT_AVAILABLE:
        return None

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()

    # Crear estilos personalizados
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        alignment=1,  # Centrado
        spaceAfter=30
    )

    header_style = ParagraphStyle(
        'CustomHeader',
        parent=styles['Heading2'],
        fontSize=16,
        spaceAfter=12
    )

    content_style = ParagraphStyle(
        'CustomContent',
        parent=styles['Normal'],
        fontSize=12,
        spaceAfter=12,
        alignment=0  # Justificado
    )

    metadata_style = ParagraphStyle(
        'Metadata',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=6
    )

    # Contenido del PDF
    story = []

    # Título
    story.append(Paragraph("Once Noticias", title_style))
    story.append(Paragraph("Sistema Editorial Optimizado", header_style))
    story.append(Spacer(1, 20))

    # Metadatos
    story.append(Paragraph("Información del Contenido", header_style))
    story.append(Paragraph(f"<b>Categoría:</b> {metadata.get('category', '')} / {metadata.get('subcategory', '')}", metadata_style))
    story.append(Paragraph(f"<b>Tipo:</b> {metadata.get('text_type', '')}", metadata_style))
    story.append(Paragraph(f"<b>Fecha:</b> {datetime.now().strftime('%d/%m/%Y %H:%M')}", metadata_style))
    story.append(Paragraph(f"<b>Tokens:</b> {metadata.get('token_count', 0):,}", metadata_style))
    story.append(Spacer(1, 30))

    # Contenido principal
    story.append(Paragraph("Contenido Generado", header_style))
    story.append(Spacer(1, 12))

    # Agregar contenido párrafo por párrafo
    for paragraph in content.split('\n\n'):
        if paragraph.strip():
            story.append(Paragraph(paragraph.strip(), content_style))

    # Pie de página
    story.append(Spacer(1, 50))
    story.append(Paragraph(f"Generado por Once Noticias AI - {datetime.now().strftime('%d/%m/%Y')}", metadata_style))

    # Construir PDF
    doc.build(story)
    buffer.seek(0)
    return buffer

def main():
    """Función principal de la aplicación"""

    # Inicializar sistemas
    prompt_system, qa_system = initialize_system()

    # Título principal
    st.title("📰 Once Noticias - Sistema Editorial Optimizado")
    st.markdown("*Reducción de 60% en tokens • Mayor velocidad • Calidad mejorada*")

    # Configuración del contenido
    col1, col2 = st.columns([2, 1])

    with col1:
        st.subheader("🎯 Configuración del Contenido")

        # Selección de parámetros
        category = st.selectbox(
            "Categoría",
            config.CATEGORIES,
            index=1  # Economía por defecto
        )

        subcategory = st.selectbox(
            "Subcategoría",
            config.SUBCATEGORIES,
            index=1  # Finanzas por defecto
        )

        text_type = st.selectbox(
            "Tipo de Texto",
            config.TEXT_TYPES,
            index=0  # Nota Periodística por defecto
        )

        # Selección de longitud
        length_option = st.selectbox(
            "Longitud del Contenido",
            list(config.LENGTH_OPTIONS.keys()),
            index=0  # Auto por defecto
        )
        selected_length = config.LENGTH_OPTIONS[length_option]

    with col2:
        st.subheader("⚙️ Configuración Avanzada")

        # Control de búsqueda web
        web_search_mode = st.selectbox(
            "🔍 Búsqueda Web",
            ["🤖 Auto (recomendado)", "🌐 Siempre activada", "📚 Solo conocimiento base"],
            index=0,
            help="Auto: usa web search cuando sea necesario. Siempre: búsqueda web en cada generación. Solo base: sin búsqueda web."
        )

        # Guardar configuración en session_state para persistencia
        st.session_state.web_search_mode = web_search_mode

        # Mostrar configuración de almacenamiento
        storage_info = prompt_system.get_storage_status()
        st.info(f"📁 {storage_info['status']}")

        # Opción para cambiar modo de almacenamiento
        if st.button("🔄 Cambiar Almacenamiento"):
            st.session_state.show_storage_config = not st.session_state.get('show_storage_config', False)

    # Configuración de almacenamiento expandible
    if st.session_state.get('show_storage_config', False):
        with st.expander("🗄️ Configuración de Almacenamiento", expanded=True):
            storage_mode = st.radio(
                "Modo de Almacenamiento",
                ["Solo Local", "Solo Snowflake (si está configurado)", "Ambos"],
                index=0
            )

            if storage_mode == "Solo Snowflake (si está configurado)":
                if not config.is_database_available():
                    st.warning("⚠️ Snowflake no está configurado. Se usará almacenamiento local.")
                else:
                    st.success("✅ Snowflake disponible")

            st.info("💡 El almacenamiento local siempre está disponible como respaldo")

    # Usar formulario para Ctrl+Enter en input principal
    st.subheader("✍️ Solicitud Editorial")

    with st.form("content_generation_form", clear_on_submit=False):
        user_prompt = st.text_area(
            "Describe el contenido que necesitas generar:",
            height=100,
            placeholder="Ejemplo: Análisis del impacto de las nuevas políticas comerciales en el sector energético mexicano...",
            help="💡 Tip: Usa Ctrl+Enter para generar contenido directamente"
        )

        # Campo para fuentes y referencias
        st.subheader("📚 Fuentes y Referencias (Opcional)")
        user_sources = st.text_area(
            "Comparte fuentes, links, datos o referencias específicas:",
            height=80,
            placeholder="Ejemplo: Reporte INEGI 2024, https://example.com/data, cifras del Banco de México, estudio de la UNAM sobre energías renovables...",
            help="Estas fuentes se integrarán al contenido junto con la búsqueda web automática"
        )

        # Botón de generación dentro del formulario
        generate_button = st.form_submit_button(
            "🚀 Generar Contenido",
            type="primary",
            use_container_width=True
        )

    # Procesamiento de la solicitud
    if generate_button and user_prompt:

        with st.spinner("🔄 Generando contenido..."):

            try:
                # Marcar tiempo de inicio
                start_time = time.time()

                # Determinar si usar Web Search basado en la selección del usuario
                force_web_search = get_web_search_setting()

                # Generar contenido con Web Search
                content_data = prompt_system.generate_content_with_web_search(
                    category=category,
                    subcategory=subcategory,
                    text_type=text_type,
                    user_prompt=user_prompt,
                    sources=user_sources,
                    selected_length=selected_length,
                    force_web_search=force_web_search
                )

                # Calcular tiempo de generación
                generation_time = time.time() - start_time

                generated_content = content_data.get("content", "")
                token_count = content_data.get("token_count", 0)
                citations = content_data.get("citations", [])
                web_search_used = content_data.get("web_search_used", False)

                # Preparar datos completos para guardar
                complete_content_data = {
                    **content_data,
                    "user_prompt": user_prompt,
                    "user_sources": user_sources,
                    "category": category,
                    "subcategory": subcategory,
                    "text_type": text_type,
                    "selected_length": selected_length,
                    "generation_time": generation_time,
                    "improvement_applied": False
                }

                # ✅ GUARDAR NUEVO CONTENIDO - crea nuevo registro
                record_id = save_new_content_metrics(prompt_system, complete_content_data, user_rating=None)

                # Guardar en session_state para flujo iterativo
                st.session_state.current_content = generated_content
                st.session_state.current_metadata = {
                    "user_prompt": user_prompt,
                    "user_sources": user_sources,
                    "category": category,
                    "subcategory": subcategory,
                    "text_type": text_type,
                    "token_count": token_count,
                    "length_setting": selected_length,
                    "citations": citations,
                    "web_search_used": web_search_used,
                    "generation_time": generation_time,
                    "record_id": record_id,  # ✅ GUARDAR EL ID DEL REGISTRO
                    "openai_call_id": content_data.get("openai_call_id", "")  # ✅ TRACK OPENAI CALL ID
                }
                st.session_state.iteration_count = 1

                # Inicializar historial de iteraciones
                st.session_state.iteration_history = {
                    1: {
                        "content": generated_content,
                        "metadata": st.session_state.current_metadata.copy(),
                        "timestamp": datetime.now().isoformat(),
                        "feedback_applied": "Contenido inicial",
                        "web_search_config": st.session_state.get('web_search_mode', '🤖 Auto (recomendado)'),
                        "record_id": record_id,  # ✅ GUARDAR EL ID EN EL HISTORIAL
                        "openai_call_id": content_data.get("openai_call_id", "")  # ✅ TRACK OPENAI CALL ID
                    }
                }
                st.session_state.current_iteration = 1

                # Mostrar solo éxito (el contenido se mostrará en la sección "Contenido Actual")
                if web_search_used:
                    st.success("✅ Contenido generado exitosamente con información actualizada de la web")
                else:
                    st.success("✅ Contenido generado exitosamente")
                st.info("📄 Ve abajo para ver tu contenido y poder mejorarlo iterativamente")

            except Exception as e:
                st.error(f"❌ Error al generar contenido: {str(e)}")
                st.error("Verifica tu configuración de OpenAI y conexión a internet")

    elif generate_button and not user_prompt:
        st.warning("⚠️ Por favor, ingresa una solicitud editorial")

    # ===== SECCIÓN DE CONTENIDO ACTUAL Y MEJORA ITERATIVA =====
    if hasattr(st.session_state, 'current_content') and st.session_state.current_content:

        # Contenedor para scroll automático
        content_container = st.container()

        with content_container:
            st.markdown("---")

            # SELECTOR DE ITERACIONES
            if hasattr(st.session_state, 'iteration_history') and len(st.session_state.iteration_history) > 1:
                col1, col2, col3 = st.columns([3, 2, 1])

                with col1:
                    st.subheader("📄 Contenido Actual")

                with col2:
                    # Selector de iteración más compacto
                    available_iterations = list(st.session_state.iteration_history.keys())
                    current_iteration = st.session_state.get('current_iteration', max(available_iterations))

                    # Formato más compacto
                    col2a, col2b = st.columns([1, 1])
                    with col2a:
                        selected_iteration = st.selectbox(
                            "Ver versión:",
                            available_iterations,
                            index=available_iterations.index(current_iteration),
                            key="iteration_selector",
                            format_func=lambda x: f"#{x}"
                        )

                    with col2b:
                        st.markdown(f"<small style='color: #666; margin-top: 24px; display: block;'>de {len(available_iterations)} total</small>", unsafe_allow_html=True)

                    # Actualizar iteración actual si cambió
                    if selected_iteration != st.session_state.get('current_iteration'):
                        st.session_state.current_iteration = selected_iteration
                        # Cargar contenido de la iteración seleccionada
                        selected_data = st.session_state.iteration_history[selected_iteration]
                        st.session_state.current_content = selected_data["content"]
                        st.session_state.current_metadata = selected_data["metadata"]
                        st.rerun()

                with col3:
                    # Mostrar timestamp de la iteración actual
                    current_iter_data = st.session_state.iteration_history[selected_iteration]
                    timestamp = current_iter_data.get("timestamp", "")
                    if timestamp:
                        try:
                            dt = datetime.fromisoformat(timestamp)
                            time_str = dt.strftime("%H:%M")
                            st.markdown(f"<small style='color: #666; margin-top: 24px; display: block;'>{time_str}</small>", unsafe_allow_html=True)
                        except:
                            pass

                # Mostrar información de la iteración seleccionada
                current_iter_data = st.session_state.iteration_history[selected_iteration]
                feedback_info = current_iter_data.get("feedback_applied", "")
                web_config = current_iter_data.get("web_search_config", "")

                if feedback_info and feedback_info != "Contenido inicial":
                    display_info = f"💬 **Mejora aplicada:** {feedback_info[:100]}..."
                    if web_config:
                        mode_short = {
                            '🤖 Auto (recomendado)': 'Auto',
                            '🌐 Siempre activada': 'Siempre Web',
                            '📚 Solo conocimiento base': 'Solo Base'
                        }.get(web_config, 'Auto')
                        display_info += f" | **Config:** {mode_short}"
                    st.info(display_info)

            else:
                # Título y iteración en la misma línea (caso original)
                col1, col2 = st.columns([4, 1])
                with col1:
                    st.subheader("📄 Contenido Actual")
                with col2:
                    st.metric("Iteración", f"#{st.session_state.get('iteration_count', 1)}", label_visibility="collapsed")

            # Mostrar información del contenido
            current_metadata = st.session_state.current_metadata
            info_text = f"**{current_metadata.get('category', '')} / {current_metadata.get('subcategory', '')}** • **{current_metadata.get('text_type', '')}**"

            # Agregar indicador de Web Search si se usó
            if current_metadata.get('web_search_used', False):
                num_sources = len(current_metadata.get('citations', []))
                if num_sources > 0:
                    info_text += f" • 🌐 **Información actualizada ({num_sources} fuentes)**"
                else:
                    info_text += " • 🌐 **Con búsqueda web**"
            else:
                info_text += " • 📚 **Solo conocimiento base**"

            # Agregar indicador de fuentes del usuario
            user_sources = current_metadata.get('user_sources', '')
            if user_sources and user_sources.strip():
                info_text += " • 📝 **Con fuentes del usuario**"

            # Agregar configuración actual de búsqueda web
            current_web_mode = st.session_state.get('web_search_mode', '🤖 Auto (recomendado)')
            mode_short = {
                '🤖 Auto (recomendado)': 'Auto',
                '🌐 Siempre activada': 'Siempre Web',
                '📚 Solo conocimiento base': 'Solo Base'
            }.get(current_web_mode, 'Auto')
            info_text += f" • **Config: {mode_short}**"

            st.markdown(info_text)

            st.markdown("---")

            # Mostrar contenido actual con citaciones
            citations = current_metadata.get('citations', [])
            display_content_with_citations(st.session_state.current_content, citations, user_sources)

            # Botones de descarga con rating de 5 estrellas
            st.markdown("---")
            col1, col2, col3 = st.columns([1, 1, 1])

            iteration = st.session_state.get('iteration_count', 1)

            # ✅ Rating de 5 estrellas - ACTUALIZA registro existente
            with col1:
                st.markdown("**⭐ Califica este contenido:**")

                # Obtener el OpenAI call ID del contenido actual
                current_openai_call_id = current_metadata.get('openai_call_id', '')

                # ✅ NEW: Clickable star rating system
                # Get current rating from session state (None by default)
                current_rating = st.session_state.get(f'content_rating_{iteration}', None)

                # Create 5 clickable stars with better styling
                star_cols = st.columns([1, 1, 1, 1, 1])
                new_rating = None

                # Add custom CSS for better star appearance
                st.markdown("""
                <style>
                .star-button {
                    font-size: 24px !important;
                    border: none !important;
                    background: transparent !important;
                    padding: 2px !important;
                    margin: 0 !important;
                }
                .star-button:hover {
                    transform: scale(1.1) !important;
                    transition: transform 0.1s !important;
                }
                </style>
                """, unsafe_allow_html=True)

                for i in range(1, 6):
                    with star_cols[i-1]:
                        # Determine star appearance
                        if current_rating is not None and i <= current_rating:
                            star_display = "⭐"  # Filled star
                            button_type = "primary"
                        else:
                            star_display = "☆"   # Empty star
                            button_type = "secondary"

                        # Create clickable star button
                        if st.button(
                            star_display,
                            key=f"star_{iteration}_{i}",
                            help=f"Calificar con {i} estrella{'s' if i > 1 else ''}",
                            use_container_width=True,
                            type=button_type
                        ):
                            new_rating = i

                # Handle rating changes
                if new_rating is not None:
                    # Update session state
                    st.session_state[f'content_rating_{iteration}'] = new_rating

                    # Update database using OpenAI call ID (multi-user safe)
                    if current_openai_call_id:
                        success = update_content_rating_by_openai_id(current_openai_call_id, new_rating)
                        if success:
                            st.success(f"✅ Rating actualizado: {new_rating}/5 estrellas")
                            st.rerun()  # Refresh to show updated stars
                        else:
                            st.error("⚠️ Error actualizando rating")
                    else:
                        st.warning("⚠️ No se puede actualizar rating - falta OpenAI call ID")

            # Word
            with col2:
                if EXPORT_AVAILABLE:
                    word_doc = create_word_document(st.session_state.current_content, current_metadata)
                    if word_doc:
                        st.download_button(
                            label="📄 Descargar Word",
                            data=word_doc,
                            file_name=f"contenido_{current_metadata.get('category', '')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"download_word_iter_{iteration}",
                            use_container_width=True
                        )
                else:
                    st.button("📄 Word (No disponible)", disabled=True, key=f"word_disabled_iter_{iteration}", use_container_width=True)

            # PDF
            with col3:
                if EXPORT_AVAILABLE:
                    pdf_doc = create_pdf_document(st.session_state.current_content, current_metadata)
                    if pdf_doc:
                        st.download_button(
                            label="📑 Descargar PDF",
                            data=pdf_doc,
                            file_name=f"contenido_{current_metadata.get('category', '')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                            mime="application/pdf",
                            key=f"download_pdf_iter_{iteration}",
                            use_container_width=True
                        )
                else:
                    st.button("📑 PDF (No disponible)", disabled=True, key=f"pdf_disabled_iter_{iteration}", use_container_width=True)

        # ===== SECCIÓN DE FEEDBACK (SOLO SI NO ESTÁ PROCESANDO) =====
        if not st.session_state.get('processing_improvement', False):
            st.markdown("---")
            st.subheader("💬 ¿Qué te pareció la nota? ¿Quisieras realizar algún ajuste?")

            # Usar formulario para Ctrl+Enter en feedback
            with st.form("feedback_form", clear_on_submit=False):
                feedback_key = f"feedback_iteration_{st.session_state.get('iteration_count', 1)}"
                user_feedback = st.text_area(
                    "Comparte tu feedback específico:",
                    height=100,
                    placeholder="Ejemplo: El título podría ser más llamativo, agregar más datos de INEGI sobre el impacto económico, incluir cifras específicas de turistas, mencionar destinos específicos como Tulum o Puerto Escondido...",
                    key=feedback_key,
                    help="💡 Tip: Usa Ctrl+Enter para aplicar mejoras directamente"
                )

                # Información sobre consistencia editorial
                st.info("💡 **Nota**: Las mejoras mantendrán automáticamente el tono, estilo y lineamientos editoriales de Once Noticias mientras aplican tu feedback.")

                # Botones dentro del formulario
                col1, col2, col3 = st.columns([1, 1, 1])

                with col1:
                    improve_button = st.form_submit_button(
                        "🚀 Mejorar Contenido",
                        type="primary",
                        use_container_width=True
                    )

                with col3:
                    new_content_button = st.form_submit_button(
                        "🗑️ Comenzar Nuevo Contenido",
                        use_container_width=True
                    )

            # Procesar botones fuera del formulario
            if improve_button:
                if user_feedback:
                    # Marcar que estamos procesando para ocultar feedback
                    st.session_state.processing_improvement = True
                    st.rerun()
                else:
                    st.warning("⚠️ Por favor, proporciona feedback específico para mejorar el contenido")

            if new_content_button:
                # Limpiar session_state
                for key in ['current_content', 'current_metadata', 'iteration_count', 'processing_improvement', 'iteration_history', 'current_iteration']:
                    if key in st.session_state:
                        del st.session_state[key]
                st.rerun()

        else:
            # Mostrar progreso y procesar mejora
            st.info("🔄 Mejorando contenido con tu feedback...")

            # Obtener el feedback del estado anterior
            feedback_key = f"feedback_iteration_{st.session_state.get('iteration_count', 1)}"
            user_feedback = st.session_state.get(feedback_key, "")

            if user_feedback:
                try:
                    # Marcar tiempo de inicio
                    start_time = time.time()

                    # Obtener configuración de búsqueda web actual (usar la del selectbox en la UI)
                    force_web_search = get_web_search_setting()

                    # Usar el nuevo método de mejora que mantiene lineamientos Once Noticias
                    content_data = prompt_system.generate_improvement_with_web_search(
                        category=current_metadata.get('category', ''),
                        subcategory=current_metadata.get('subcategory', ''),
                        text_type=current_metadata.get('text_type', ''),
                        initial_content=st.session_state.current_content,
                        user_feedback=user_feedback,
                        sources=current_metadata.get('user_sources', ''),
                        selected_length=current_metadata.get('length_setting', 'auto'),
                        force_web_search=force_web_search
                    )

                    # Calcular tiempo de generación
                    generation_time = time.time() - start_time

                    improved_content = content_data.get("content", "")
                    improved_tokens = content_data.get("token_count", 0)
                    improved_citations = content_data.get("citations", [])
                    improved_web_search_used = content_data.get("web_search_used", False)

                    # Preparar datos completos para guardar
                    improvement_data = {
                        **content_data,
                        "user_prompt": current_metadata.get('user_prompt', ''),
                        "user_sources": current_metadata.get('user_sources', ''),
                        "user_feedback": user_feedback,
                        "category": current_metadata.get('category', ''),
                        "subcategory": current_metadata.get('subcategory', ''),
                        "text_type": current_metadata.get('text_type', ''),
                        "selected_length": current_metadata.get('length_setting', 'auto'),
                        "generation_time": generation_time,
                        "improvement_applied": True
                    }

                    # ✅ GUARDAR CONTENIDO MEJORADO - crea NUEVO registro
                    new_record_id = save_new_content_metrics(prompt_system, improvement_data, user_rating=None)

                    # Actualizar session_state con el contenido mejorado
                    st.session_state.current_content = improved_content
                    st.session_state.current_metadata.update({
                        "token_count": improved_tokens,
                        "last_feedback": user_feedback,
                        "citations": improved_citations,
                        "web_search_used": improved_web_search_used,
                        "record_id": new_record_id,  # ✅ NUEVO ID PARA EL CONTENIDO MEJORADO
                        "openai_call_id": content_data.get("openai_call_id", "")  # ✅ NUEVO OPENAI CALL ID
                    })
                    st.session_state.iteration_count += 1
                    st.session_state.processing_improvement = False

                    # Guardar nueva iteración en historial
                    new_iteration = st.session_state.iteration_count
                    st.session_state.iteration_history[new_iteration] = {
                        "content": improved_content,
                        "metadata": st.session_state.current_metadata.copy(),
                        "timestamp": datetime.now().isoformat(),
                        "feedback_applied": user_feedback[:200],  # Primeros 200 chars del feedback
                        "web_search_config": st.session_state.get('web_search_mode', '🤖 Auto (recomendado)'),
                        "record_id": new_record_id,  # ✅ NUEVO ID EN EL HISTORIAL
                        "openai_call_id": content_data.get("openai_call_id", "")  # ✅ TRACK OPENAI CALL ID
                    }
                    st.session_state.current_iteration = new_iteration

                    if improved_web_search_used:
                        st.success("✅ Contenido mejorado exitosamente con información actualizada")
                    else:
                        st.success("✅ Contenido mejorado exitosamente")

                    st.rerun()

                except Exception as e:
                    st.session_state.processing_improvement = False
                    st.error(f"❌ Error al mejorar contenido: {str(e)}")
                    st.rerun()

if __name__ == "__main__":
    main()